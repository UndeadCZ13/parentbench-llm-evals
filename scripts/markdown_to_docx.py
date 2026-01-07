#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Convert Markdown report to DOCX with:
- Proper heading levels
- Embedded images (relative paths)
- Clean paragraph structure
"""

from pathlib import Path
import re

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import Pt


IMAGE_PATTERN = re.compile(r"!\[(.*?)\]\((.*?)\)")


def add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    return h


def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    return p


def add_image(doc: Document, image_path: Path, max_width_inch: float = 6.0):
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(max_width_inch))
    else:
        doc.add_paragraph(f"[Missing image: {image_path}]")


def markdown_to_docx(md_path: Path, out_path: Path):
    doc = Document()

    # Optional: set default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    for line in lines:
        line = line.rstrip()

        # Headings
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), level=1)
            continue
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
            continue
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=3)
            continue

        # Images
        img_match = IMAGE_PATTERN.search(line)
        if img_match:
            alt, path = img_match.groups()
            img_path = (md_path.parent / path).resolve()
            add_image(doc, img_path)
            if alt:
                cap = doc.add_paragraph(f"Figure: {alt}")
                cap.italic = True
            continue

        # Empty line
        if not line.strip():
            doc.add_paragraph("")
            continue

        # Normal paragraph
        add_paragraph(doc, line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"[OK] DOCX written to: {out_path}")


def main():
    import argparse

    ap = argparse.ArgumentParser()
    ap.add_argument("--md", required=True, help="Input markdown file")
    ap.add_argument("--out", required=True, help="Output docx file")
    args = ap.parse_args()

    markdown_to_docx(Path(args.md), Path(args.out))


if __name__ == "__main__":
    main()
